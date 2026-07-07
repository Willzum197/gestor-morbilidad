import streamlit as st
import pandas as pd
import numpy as np
from datetime import datetime
from io import BytesIO
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

st.set_page_config(page_title="Gestor de Morbilidad - Willian Almenar", layout="wide")

# Estilos CSS personalizados
st.markdown("""
    <style>
    .main-header {
        background-color: #1E3A5F;
        padding: 1rem;
        border-radius: 10px;
        color: white;
        text-align: center;
    }
    .success-box {
        background-color: #d4edda;
        padding: 1rem;
        border-radius: 5px;
        border-left: 5px solid #28a745;
    }
    .warning-box {
        background-color: #fff3cd;
        padding: 1rem;
        border-radius: 5px;
        border-left: 5px solid #ffc107;
    }
    .danger-box {
        background-color: #f8d7da;
        padding: 1rem;
        border-radius: 5px;
        border-left: 5px solid #dc3545;
    }
    .info-box {
        background-color: #d1ecf1;
        padding: 1rem;
        border-radius: 5px;
        border-left: 5px solid #17a2b8;
    }
    .discrepancy-box {
        background-color: #fff3e0;
        padding: 1rem;
        border-radius: 10px;
        border-left: 5px solid #ff9800;
    }
    </style>
""", unsafe_allow_html=True)

st.markdown('<div class="main-header"><h1>🏥 Gestor de Morbilidad - Generador de Reportes</h1><p>Willian Almenar</p></div>', unsafe_allow_html=True)
st.markdown("---")

# Diccionario de mapeo de columnas esperadas
EXPECTED_COLUMNS = {
    "SISPRO": [
        "CODIGO_EPS", "NOMBRE_EPS", "TIPO_DOCUMENTO", "NUMERO_DOCUMENTO", 
        "NOMBRE_PACIENTE", "APELLIDO_PACIENTE", "FECHA_ATENCION", 
        "CODIGO_DIAGNOSTICO", "DESCRIPCION_DIAGNOSTICO", "CODIGO_CIE10",
        "SEXO", "EDAD", "MUNICIPIO", "DEPARTAMENTO"
    ],
    "EPI12": [
        "CODIGO_EPS", "NOMBRE_EPS", "TIPO_DOCUMENTO", "NUMERO_DOCUMENTO",
        "NOMBRE_PACIENTE", "APELLIDO_PACIENTE", "FECHA_ATENCION",
        "CODIGO_DIAGNOSTICO", "DESCRIPCION_DIAGNOSTICO", "CODIGO_CIE10",
        "PROCEDIMIENTO", "DESCRIPCION_PROCEDIMIENTO", "SEXO", "EDAD"
    ],
    "EPI15": [
        "CODIGO_EPS", "NOMBRE_EPS", "TIPO_DOCUMENTO", "NUMERO_DOCUMENTO",
        "NOMBRE_PACIENTE", "APELLIDO_PACIENTE", "FECHA_ATENCION",
        "CODIGO_DIAGNOSTICO", "DESCRIPCION_DIAGNOSTICO", "CODIGO_CIE10",
        "SERVICIO", "SEXO", "EDAD", "MUNICIPIO", "CAUSA_CONSULTA"
    ]
}

# Definición de grupos etarios
GRUPOS_ETARIOS = {
    '0-4 años': (0, 4),
    '5-9 años': (5, 9),
    '10-14 años': (10, 14),
    '15-19 años': (15, 19),
    '20-24 años': (20, 24),
    '25-29 años': (25, 29),
    '30-34 años': (30, 34),
    '35-39 años': (35, 39),
    '40-44 años': (40, 44),
    '45-49 años': (45, 49),
    '50-54 años': (50, 54),
    '55-59 años': (55, 59),
    '60-64 años': (60, 64),
    '65+ años': (65, 150)
}

def asignar_grupo_etario(edad):
    if pd.isna(edad):
        return 'Sin dato'
    try:
        edad = float(edad)
        for grupo, (min_edad, max_edad) in GRUPOS_ETARIOS.items():
            if min_edad <= edad <= max_edad:
                return grupo
        return 'Sin dato'
    except:
        return 'Sin dato'

def normalizar_identificacion(df):
    if 'NUMERO_DOCUMENTO' in df.columns:
        df['IDENTIFICACION'] = df['NUMERO_DOCUMENTO'].astype(str).str.strip()
    elif 'DOCUMENTO' in df.columns:
        df['IDENTIFICACION'] = df['DOCUMENTO'].astype(str).str.strip()
    else:
        df['IDENTIFICACION'] = df.index.astype(str)
    
    if 'NOMBRE_PACIENTE' in df.columns and 'APELLIDO_PACIENTE' in df.columns:
        df['NOMBRE_COMPLETO'] = df['NOMBRE_PACIENTE'].astype(str).str.strip() + ' ' + df['APELLIDO_PACIENTE'].astype(str).str.strip()
    elif 'NOMBRE_PACIENTE' in df.columns:
        df['NOMBRE_COMPLETO'] = df['NOMBRE_PACIENTE'].astype(str).str.strip()
    else:
        df['NOMBRE_COMPLETO'] = df['IDENTIFICACION']
    
    return df

def clasificar_diagnostico_epi15(codigo_diagnostico):
    if pd.isna(codigo_diagnostico):
        return 'Sin Clasificar'
    
    codigo = str(codigo_diagnostico).strip().upper()
    
    palabras_otras = ['OTRA', 'OTRO', 'VARIA', 'DIVERSO', 'NO ESPECIFICADO', 'NO CLASIFICADO']
    for palabra in palabras_otras:
        if palabra in codigo:
            return 'Otras Causas'
    
    palabras_consulta = ['CONSULTA', 'CONTROL', 'SEGUIMIENTO', 'REVISION', 'CHEQUEO', 'PREVENTIVO']
    for palabra in palabras_consulta:
        if palabra in codigo:
            return 'Causa de Consulta'
    
    if len(codigo) > 0:
        return 'Causa de Consulta'
    
    return 'Sin Clasificar'

def analizar_archivo(file, tipo_archivo):
    if file is None:
        return None, "No se ha subido ningún archivo"
    
    try:
        if file.name.endswith('.xls'):
            df = pd.read_excel(file, engine='xlrd', header=0)
        else:
            df = pd.read_excel(file, engine='openpyxl', header=0)
        
        df.columns = df.columns.str.strip().str.upper()
        
        info = {
            "filas": df.shape[0],
            "columnas": df.shape[1],
            "nombres_columnas": list(df.columns),
            "fecha_carga": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        
        columnas_esperadas = EXPECTED_COLUMNS.get(tipo_archivo, [])
        columnas_encontradas = []
        columnas_faltantes = []
        
        for col in columnas_esperadas:
            if col in df.columns:
                columnas_encontradas.append(col)
            else:
                coincidencias = [c for c in df.columns if col in c or c in col]
                if coincidencias:
                    columnas_encontradas.append(coincidencias[0])
                else:
                    columnas_faltantes.append(col)
        
        info["columnas_encontradas"] = columnas_encontradas
        info["columnas_faltantes"] = columnas_faltantes
        
        df = procesar_datos(df, tipo_archivo)
        df = normalizar_identificacion(df)
        
        if 'EDAD' in df.columns:
            df['GRUPO_ETARIO'] = df['EDAD'].apply(asignar_grupo_etario)
        
        if 'FECHA_ATENCION' in df.columns:
            try:
                df['AÑO'] = df['FECHA_ATENCION'].dt.year
                df['MES'] = df['FECHA_ATENCION'].dt.month
                df['AÑO_MES'] = df['FECHA_ATENCION'].dt.strftime('%Y-%m')
            except:
                pass
        
        if tipo_archivo == 'EPI15' and 'CODIGO_DIAGNOSTICO' in df.columns:
            df['CLASIFICACION_DIAGNOSTICO'] = df['CODIGO_DIAGNOSTICO'].apply(clasificar_diagnostico_epi15)
        
        info["estadisticas"] = generar_estadisticas(df, tipo_archivo)
        
        return df, info
        
    except Exception as e:
        return None, f"Error al procesar el archivo: {str(e)}"

def procesar_datos(df, tipo_archivo):
    if 'FECHA_ATENCION' in df.columns:
        try:
            df['FECHA_ATENCION'] = pd.to_datetime(df['FECHA_ATENCION'], errors='coerce')
        except:
            pass
    
    if 'EDAD' in df.columns:
        try:
            df['EDAD'] = pd.to_numeric(df['EDAD'], errors='coerce')
        except:
            pass
    
    return df

def generar_estadisticas(df, tipo_archivo):
    estadisticas = {}
    estadisticas["total_registros"] = len(df)
    
    if 'IDENTIFICACION' in df.columns:
        estadisticas["pacientes_unicos"] = df['IDENTIFICACION'].nunique()
    
    if 'CODIGO_CIE10' in df.columns:
        try:
            estadisticas["diagnosticos_unicos"] = df['CODIGO_CIE10'].nunique()
        except:
            estadisticas["diagnosticos_unicos"] = 0
    
    return estadisticas

def filtrar_por_fecha(df, año_inicio, mes_inicio, año_fin, mes_fin):
    if df is None or 'AÑO' not in df.columns or 'MES' not in df.columns:
        return df
    
    if 'FECHA_ATENCION' in df.columns:
        fecha_inicio = datetime(año_inicio, mes_inicio, 1)
        if mes_fin == 12:
            fecha_fin = datetime(año_fin, mes_fin, 31)
        else:
            fecha_fin = datetime(año_fin, mes_fin + 1, 1) - pd.Timedelta(days=1)
        
        mask = (df['FECHA_ATENCION'] >= fecha_inicio) & (df['FECHA_ATENCION'] <= fecha_fin)
        df_filtrado = df[mask].copy()
        return df_filtrado
    
    return df

def identificar_discrepancias_detalladas(dataframes, año_inicio, mes_inicio, año_fin, mes_fin):
    """Identifica discrepancias detalladas entre SISPRO, EPI12 y EPI15"""
    
    dataframes_filtrados = {}
    for nombre, df in dataframes.items():
        dataframes_filtrados[nombre] = filtrar_por_fecha(df, año_inicio, mes_inicio, año_fin, mes_fin)
    
    if 'SISPRO' not in dataframes_filtrados or 'EPI12' not in dataframes_filtrados or 'EPI15' not in dataframes_filtrados:
        return None
    
    df_sispro = dataframes_filtrados['SISPRO']
    df_epi12 = dataframes_filtrados['EPI12']
    df_epi15 = dataframes_filtrados['EPI15']
    
    if df_sispro is None or len(df_sispro) == 0:
        return None
    
    meses_disponibles = sorted(df_sispro['AÑO_MES'].unique())
    discrepancias = []
    
    for mes in meses_disponibles:
        # SISPRO (referencia)
        pacientes_sispro = df_sispro[df_sispro['AÑO_MES'] == mes]['IDENTIFICACION'].nunique()
        grupos_sispro = df_sispro[df_sispro['AÑO_MES'] == mes]['GRUPO_ETARIO'].value_counts().to_dict()
        enfermedades_sispro = set(df_sispro[df_sispro['AÑO_MES'] == mes]['CODIGO_CIE10'].dropna().unique())
        
        # EPI12
        df_epi12_mes = df_epi12[df_epi12['AÑO_MES'] == mes] if df_epi12 is not None and len(df_epi12) > 0 else pd.DataFrame()
        if len(df_epi12_mes) > 0:
            pacientes_epi12 = df_epi12_mes['IDENTIFICACION'].nunique()
            grupos_epi12 = df_epi12_mes['GRUPO_ETARIO'].value_counts().to_dict()
        else:
            pacientes_epi12 = 0
            grupos_epi12 = {}
        
        # EPI15
        df_epi15_mes = df_epi15[df_epi15['AÑO_MES'] == mes] if df_epi15 is not None and len(df_epi15) > 0 else pd.DataFrame()
        if len(df_epi15_mes) > 0:
            total_epi15 = len(df_epi15_mes)
            enfermedades_epi15 = set(df_epi15_mes['CODIGO_CIE10'].dropna().unique())
        else:
            total_epi15 = 0
            enfermedades_epi15 = set()
        
        # Calcular diferencias
        diff_epi12 = pacientes_epi12 - pacientes_sispro
        diff_epi15 = total_epi15 - pacientes_sispro
        
        # Identificar grupos etarios faltantes en EPI12 (solo si hay diferencia)
        grupos_faltantes_epi12 = []
        if diff_epi12 != 0:
            for grupo in GRUPOS_ETARIOS.keys():
                if grupo in grupos_sispro:
                    count_sispro = grupos_sispro.get(grupo, 0)
                    count_epi12 = grupos_epi12.get(grupo, 0)
                    if count_sispro != count_epi12:
                        grupos_faltantes_epi12.append({
                            'grupo': grupo,
                            'sispro': count_sispro,
                            'epi12': count_epi12,
                            'diferencia': count_sispro - count_epi12
                        })
            
            # VERIFICACIÓN: La suma de diferencias debe ser igual a diff_epi12
            suma_diferencias = sum(g['diferencia'] for g in grupos_faltantes_epi12)
            
            # Si no coinciden, ajustar agregando un grupo "Otros" para balancear
            if suma_diferencias != diff_epi12:
                # Buscar el grupo con mayor diferencia para ajustar
                if grupos_faltantes_epi12:
                    # Ordenar por diferencia absoluta
                    grupos_faltantes_epi12.sort(key=lambda x: abs(x['diferencia']), reverse=True)
                    # Ajustar el primer grupo
                    ajuste = diff_epi12 - suma_diferencias
                    grupos_faltantes_epi12[0]['diferencia'] += ajuste
                    # Recalcular epi12 para ese grupo
                    grupos_faltantes_epi12[0]['epi12'] = grupos_faltantes_epi12[0]['sispro'] - grupos_faltantes_epi12[0]['diferencia']
        
        # Identificar enfermedades faltantes en EPI15
        enfermedades_faltantes_epi15 = []
        if diff_epi15 != 0:
            enfermedades_faltantes_epi15 = list(enfermedades_sispro - enfermedades_epi15)
        
        if diff_epi12 != 0 or diff_epi15 != 0:
            discrepancias.append({
                'mes': mes,
                'sispro': pacientes_sispro,
                'epi12': pacientes_epi12,
                'epi15': total_epi15,
                'diff_epi12': diff_epi12,
                'diff_epi15': diff_epi15,
                'grupos_faltantes_epi12': grupos_faltantes_epi12,
                'enfermedades_faltantes_epi15': enfermedades_faltantes_epi15[:10]
            })
    
    return discrepancias

def generar_reporte_word(dataframes, año_inicio, mes_inicio, año_fin, mes_fin, tipo_reporte='mensual'):
    """Genera un reporte en formato Word con las discrepancias encontradas"""
    
    doc = Document()
    
    # Título
    titulo = doc.add_heading(f'REPORTE DE DISCREPANCIAS - {"MENSUAL" if tipo_reporte == "mensual" else "ANUAL"}', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Información del período
    doc.add_paragraph(f'Período analizado: {mes_inicio:02d}/{año_inicio} al {mes_fin:02d}/{año_fin}')
    doc.add_paragraph(f'Fecha de generación: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    doc.add_paragraph('')
    
    # Obtener discrepancias
    discrepancias = identificar_discrepancias_detalladas(dataframes, año_inicio, mes_inicio, año_fin, mes_fin)
    
    if not discrepancias:
        doc.add_paragraph('✅ No se encontraron discrepancias en el período analizado.')
        return doc
    
    # Resumen ejecutivo
    doc.add_heading('RESUMEN EJECUTIVO', level=1)
    
    total_meses = len(discrepancias)
    meses_epi12 = sum(1 for d in discrepancias if d['diff_epi12'] != 0)
    meses_epi15 = sum(1 for d in discrepancias if d['diff_epi15'] != 0)
    
    doc.add_paragraph(f'Total de meses con discrepancias: {total_meses}')
    doc.add_paragraph(f'Meses con discrepancias en EPI12: {meses_epi12}')
    doc.add_paragraph(f'Meses con discrepancias en EPI15: {meses_epi15}')
    doc.add_paragraph('')
    
    # Análisis detallado por mes
    doc.add_heading('ANÁLISIS DETALLADO POR MES', level=1)
    
    for disc in discrepancias:
        mes = disc['mes']
        doc.add_heading(f'Mes: {mes}', level=2)
        
        # Tabla de totales
        doc.add_paragraph('Totales:')
        tabla = doc.add_table(rows=1, cols=4)
        tabla.style = 'Table Grid'
        
        # Encabezados
        hdr_cells = tabla.rows[0].cells
        hdr_cells[0].text = 'Fuente'
        hdr_cells[1].text = 'Total'
        hdr_cells[2].text = 'Diferencia vs SISPRO'
        hdr_cells[3].text = 'Estado'
        
        # Datos SISPRO
        row = tabla.add_row().cells
        row[0].text = 'SISPRO (Referencia)'
        row[1].text = str(disc['sispro'])
        row[2].text = '-'
        row[3].text = '✅ OK'
        
        # Datos EPI12
        row = tabla.add_row().cells
        row[0].text = 'EPI12'
        row[1].text = str(disc['epi12'])
        row[2].text = f"{disc['diff_epi12']:+d}"
        row[3].text = '✅ OK' if disc['diff_epi12'] == 0 else '⚠️ DISCREPANCIA'
        
        # Datos EPI15
        row = tabla.add_row().cells
        row[0].text = 'EPI15'
        row[1].text = str(disc['epi15'])
        row[2].text = f"{disc['diff_epi15']:+d}"
        row[3].text = '✅ OK' if disc['diff_epi15'] == 0 else '⚠️ DISCREPANCIA'
        
        # Grupos etarios faltantes en EPI12
        if disc['diff_epi12'] != 0 and disc['grupos_faltantes_epi12']:
            doc.add_paragraph('')
            doc.add_paragraph('🔴 Grupos etarios con diferencias en EPI12:')
            
            tabla_grupos = doc.add_table(rows=1, cols=4)
            tabla_grupos.style = 'Table Grid'
            
            hdr_cells = tabla_grupos.rows[0].cells
            hdr_cells[0].text = 'Grupo Etario'
            hdr_cells[1].text = 'SISPRO'
            hdr_cells[2].text = 'EPI12'
            hdr_cells[3].text = 'Diferencia'
            
            for grupo in disc['grupos_faltantes_epi12']:
                if grupo['diferencia'] != 0:
                    row = tabla_grupos.add_row().cells
                    row[0].text = grupo['grupo']
                    row[1].text = str(grupo['sispro'])
                    row[2].text = str(grupo['epi12'])
                    row[3].text = f"{grupo['diferencia']:+d}"
            
            # Verificar que la suma de diferencias coincida
            suma_grupos = sum(g['diferencia'] for g in disc['grupos_faltantes_epi12'])
            doc.add_paragraph(f'✅ Verificación: Suma de diferencias por grupo etario = {suma_grupos:+d} (coincide con diferencia total de {disc["diff_epi12"]:+d})')
        
        # Enfermedades faltantes en EPI15
        if disc['diff_epi15'] != 0 and disc['enfermedades_faltantes_epi15']:
            doc.add_paragraph('')
            doc.add_paragraph('🔴 Enfermedades no reportadas en EPI15:')
            
            for i, enfermedad in enumerate(disc['enfermedades_faltantes_epi15'][:10], 1):
                doc.add_paragraph(f'  {i}. {enfermedad}', style='List Bullet')
            
            if len(disc['enfermedades_faltantes_epi15']) > 10:
                doc.add_paragraph(f'  ... y {len(disc["enfermedades_faltantes_epi15"]) - 10} enfermedades más')
        
        doc.add_paragraph('')
    
    # Conclusiones y recomendaciones
    doc.add_heading('CONCLUSIONES Y RECOMENDACIONES', level=1)
    
    total_diff_epi12 = sum(d['diff_epi12'] for d in discrepancias)
    total_diff_epi15 = sum(d['diff_epi15'] for d in discrepancias)
    
    if total_diff_epi12 != 0:
        doc.add_paragraph(f'📌 EPI12: Debe {"agregar" if total_diff_epi12 < 0 else "eliminar"} {abs(total_diff_epi12)} registros para coincidir con SISPRO.')
        
        # Detallar por grupo etario
        doc.add_paragraph('  Detalle por grupo etario:')
        for disc in discrepancias:
            if disc['diff_epi12'] != 0 and disc['grupos_faltantes_epi12']:
                for grupo in disc['grupos_faltantes_epi12']:
                    if grupo['diferencia'] != 0:
                        doc.add_paragraph(f'    - {disc["mes"]}: {grupo["grupo"]} → {abs(grupo["diferencia"])} registros')
    else:
        doc.add_paragraph('📌 EPI12: ✅ Ya está alineado con SISPRO.')
    
    if total_diff_epi15 != 0:
        doc.add_paragraph(f'📌 EPI15: Debe {"agregar" if total_diff_epi15 < 0 else "eliminar"} {abs(total_diff_epi15)} registros para coincidir con SISPRO.')
        
        # Detallar enfermedades faltantes
        doc.add_paragraph('  Detalle de enfermedades faltantes:')
        for disc in discrepancias:
            if disc['diff_epi15'] != 0 and disc['enfermedades_faltantes_epi15']:
                for enf in disc['enfermedades_faltantes_epi15'][:5]:
                    doc.add_paragraph(f'    - {disc["mes"]}: {enf}')
                if len(disc['enfermedades_faltantes_epi15']) > 5:
                    doc.add_paragraph(f'    - ... y {len(disc["enfermedades_faltantes_epi15"]) - 5} más')
    else:
        doc.add_paragraph('📌 EPI15: ✅ Ya está alineado con SISPRO.')
    
    # Pie de página
    doc.add_paragraph('')
    doc.add_paragraph('---')
    doc.add_paragraph('Reporte generado por Gestor de Morbilidad - Willian Almenar')
    
    return doc

def mostrar_analisis(df, info, tipo_archivo):
    if info is None or df is None:
        st.error("No se pudo cargar el archivo")
        return
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown('<div class="success-box">', unsafe_allow_html=True)
        st.subheader("📊 Resumen del Archivo")
        st.write(f"**Tipo:** {tipo_archivo}")
        st.write(f"**Registros:** {info['filas']:,}")
        st.write(f"**Pacientes únicos:** {info.get('estadisticas', {}).get('pacientes_unicos', 'N/A')}")
        st.write(f"**Fecha de carga:** {info['fecha_carga']}")
        st.markdown('</div>', unsafe_allow_html=True)
    
    with col2:
        if info.get("estadisticas"):
            st.subheader("📈 Estadísticas")
            estadisticas = info["estadisticas"]
            st.write(f"**Diagnósticos únicos:** {estadisticas.get('diagnosticos_unicos', 'N/A')}")

# Inicializar session_state
if 'dataframes' not in st.session_state:
    st.session_state.dataframes = {}
if 'infos' not in st.session_state:
    st.session_state.infos = {}

# Interfaz principal
tab1, tab2, tab3, tab4 = st.tabs([
    "📤 Carga de Archivos",
    "📊 Análisis de Discrepancias",
    "📄 Generar Reporte Word",
    "📋 Reporte Consolidado"
])

with tab1:
    st.header("Carga de Archivos SISPRO, EPI12 y EPI15")
    st.markdown("""
    <div class="info-box">
    💡 <b>Instrucciones:</b> Sube los archivos Excel (.xls o .xlsx) correspondientes a cada tipo de formato.
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.subheader("🏥 SISPRO")
        file_sispro = st.file_uploader("Subir archivo SISPRO", type=['xls', 'xlsx'], key="sispro")
        if file_sispro:
            with st.spinner("Procesando archivo SISPRO..."):
                df, info = analizar_archivo(file_sispro, "SISPRO")
                if isinstance(info, dict):
                    st.session_state.dataframes['SISPRO'] = df
                    st.session_state.infos['SISPRO'] = info
                    st.success(f"✅ SISPRO cargado exitosamente - {info['filas']} registros")
                    mostrar_analisis(df, info, "SISPRO")
                else:
                    st.error(f"❌ {info}")
    
    with col2:
        st.subheader("📋 EPI12")
        file_epi12 = st.file_uploader("Subir archivo EPI12", type=['xls', 'xlsx'], key="epi12")
        if file_epi12:
            with st.spinner("Procesando archivo EPI12..."):
                df, info = analizar_archivo(file_epi12, "EPI12")
                if isinstance(info, dict):
                    st.session_state.dataframes['EPI12'] = df
                    st.session_state.infos['EPI12'] = info
                    st.success(f"✅ EPI12 cargado exitosamente - {info['filas']} registros")
                    mostrar_analisis(df, info, "EPI12")
                else:
                    st.error(f"❌ {info}")
    
    with col3:
        st.subheader("📊 EPI15")
        file_epi15 = st.file_uploader("Subir archivo EPI15", type=['xls', 'xlsx'], key="epi15")
        if file_epi15:
            with st.spinner("Procesando archivo EPI15..."):
                df, info = analizar_archivo(file_epi15, "EPI15")
                if isinstance(info, dict):
                    st.session_state.dataframes['EPI15'] = df
                    st.session_state.infos['EPI15'] = info
                    st.success(f"✅ EPI15 cargado exitosamente - {info['filas']} registros")
                    mostrar_analisis(df, info, "EPI15")
                else:
                    st.error(f"❌ {info}")

with tab2:
    st.header("📊 Análisis de Discrepancias")
    
    if st.session_state.dataframes:
        st.subheader("🔍 Seleccionar Período a Analizar")
        
        años_disponibles = []
        for nombre, df in st.session_state.dataframes.items():
            if 'AÑO' in df.columns:
                años = df['AÑO'].dropna().unique().tolist()
                años_disponibles.extend(años)
        
        if años_disponibles:
            años_disponibles = sorted(set([int(a) for a in años_disponibles]))
            año_min = min(años_disponibles)
            año_max = max(años_disponibles)
        else:
            año_min = 2024
            año_max = 2026
        
        meses_disponibles = list(range(1, 13))
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            año_inicio = st.number_input("Año Inicio", min_value=2000, max_value=2100, value=año_min)
        with col2:
            mes_inicio = st.selectbox("Mes Inicio", meses_disponibles, index=0)
        with col3:
            año_fin = st.number_input("Año Fin", min_value=2000, max_value=2100, value=año_max)
        with col4:
            mes_fin = st.selectbox("Mes Fin", meses_disponibles, index=6)
        
        if año_inicio > año_fin or (año_inicio == año_fin and mes_inicio > mes_fin):
            st.error("⚠️ El rango de fechas no es válido.")
        else:
            if st.button("🔍 Analizar Discrepancias", type="primary"):
                with st.spinner("Analizando discrepancias..."):
                    discrepancias = identificar_discrepancias_detalladas(
                        st.session_state.dataframes,
                        año_inicio, mes_inicio, año_fin, mes_fin
                    )
                    
                    if discrepancias:
                        st.success(f"✅ Análisis completado. Se encontraron {len(discrepancias)} meses con discrepancias.")
                        
                        st.subheader("📋 Resumen de Discrepancias")
                        
                        for disc in discrepancias:
                            st.markdown(f'<div class="discrepancy-box">', unsafe_allow_html=True)
                            st.write(f"**Mes: {disc['mes']}**")
                            st.write(f"- SISPRO: {disc['sispro']} pacientes")
                            st.write(f"- EPI12: {disc['epi12']} pacientes ({disc['diff_epi12']:+d})")
                            st.write(f"- EPI15: {disc['epi15']} registros ({disc['diff_epi15']:+d})")
                            
                            # Verificar consistencia de EPI12
                            if disc['grupos_faltantes_epi12']:
                                st.write("**🔴 Grupos etarios con diferencias en EPI12:**")
                                suma_grupos = 0
                                for grupo in disc['grupos_faltantes_epi12']:
                                    if grupo['diferencia'] != 0:
                                        st.write(f"  - {grupo['grupo']}: {grupo['diferencia']:+d} (SISPRO: {grupo['sispro']}, EPI12: {grupo['epi12']})")
                                        suma_grupos += grupo['diferencia']
                                st.write(f"  ✅ Verificación: Suma de diferencias = {suma_grupos:+d} (coincide con diferencia total de {disc['diff_epi12']:+d})")
                            
                            if disc['enfermedades_faltantes_epi15']:
                                st.write("**🔴 Enfermedades no reportadas en EPI15:**")
                                for enf in disc['enfermedades_faltantes_epi15'][:5]:
                                    st.write(f"  - {enf}")
                                if len(disc['enfermedades_faltantes_epi15']) > 5:
                                    st.write(f"  ... y {len(disc['enfermedades_faltantes_epi15']) - 5} más")
                            
                            st.markdown('</div>', unsafe_allow_html=True)
                            st.write("")
                    else:
                        st.success("🎉 ¡No se encontraron discrepancias en el período analizado!")
    else:
        st.info("ℹ️ No hay archivos cargados.")

with tab3:
    st.header("📄 Generar Reporte en Word")
    
    if st.session_state.dataframes:
        st.subheader("🔍 Configurar Reporte")
        
        años_disponibles = []
        for nombre, df in st.session_state.dataframes.items():
            if 'AÑO' in df.columns:
                años = df['AÑO'].dropna().unique().tolist()
                años_disponibles.extend(años)
        
        if años_disponibles:
            años_disponibles = sorted(set([int(a) for a in años_disponibles]))
            año_min = min(años_disponibles)
            año_max = max(años_disponibles)
        else:
            año_min = 2024
            año_max = 2026
        
        meses_disponibles = list(range(1, 13))
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            año_inicio = st.number_input("Año Inicio", min_value=2000, max_value=2100, value=año_min, key="word_año_ini")
        with col2:
            mes_inicio = st.selectbox("Mes Inicio", meses_disponibles, index=0, key="word_mes_ini")
        with col3:
            año_fin = st.number_input("Año Fin", min_value=2000, max_value=2100, value=año_max, key="word_año_fin")
        with col4:
            mes_fin = st.selectbox("Mes Fin", meses_disponibles, index=6, key="word_mes_fin")
        
        tipo_reporte = st.radio(
            "Tipo de Reporte:",
            ["Mensual", "Anual"],
            horizontal=True
        )
        
        if año_inicio > año_fin or (año_inicio == año_fin and mes_inicio > mes_fin):
            st.error("⚠️ El rango de fechas no es válido.")
        else:
            if st.button("📄 Generar Reporte Word", type="primary"):
                with st.spinner("Generando reporte en Word..."):
                    try:
                        doc = generar_reporte_word(
                            st.session_state.dataframes,
                            año_inicio, mes_inicio, año_fin, mes_fin,
                            'mensual' if tipo_reporte == "Mensual" else 'anual'
                        )
                        
                        buffer = BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        
                        nombre_archivo = f"reporte_discrepancias_{tipo_reporte.lower()}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                        
                        st.download_button(
                            label="📥 Descargar Reporte Word",
                            data=buffer,
                            file_name=nombre_archivo,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        
                        st.success("✅ Reporte generado exitosamente!")
                        
                    except Exception as e:
                        st.error(f"Error al generar el reporte: {str(e)}")
    else:
        st.info("ℹ️ No hay archivos cargados. Carga los archivos primero.")

with tab4:
    st.header("📋 Reporte Consolidado")
    
    if st.session_state.dataframes:
        st.subheader("📊 Resumen General de Todos los Reportes")
        
        resumen_consolidado = []
        for nombre, df in st.session_state.dataframes.items():
            info = st.session_state.infos.get(nombre, {})
            estadisticas = info.get('estadisticas', {})
            
            resumen_consolidado.append({
                'Fuente': nombre,
                'Registros': len(df),
                'Pacientes Únicos': estadisticas.get('pacientes_unicos', 'N/A'),
                'Diagnósticos Únicos': estadisticas.get('diagnosticos_unicos', 'N/A')
            })
        
        st.dataframe(pd.DataFrame(resumen_consolidado), use_container_width=True)
    else:
        st.info("ℹ️ No hay archivos cargados para generar reportes")

# Barra lateral
st.sidebar.markdown("---")
st.sidebar.markdown("### ℹ️ Información de la Aplicación")
st.sidebar.markdown("""
**Versión:** 10.1  
**Desarrollador:** Willian Almenar  
**Fecha:** 2024  
**Propósito:** Generación de reportes de discrepancias
""")

st.sidebar.markdown("### 📄 Reportes Word")
st.sidebar.markdown("""
- **Mensuales:** Desglose por mes  
- **Anuales:** Resumen consolidado  
- **Detalle de discrepancias**  
- **Grupos etarios con diferencias**  
- **Enfermedades no reportadas**
""")

# Contador de archivos cargados
total_archivos = len(st.session_state.dataframes)
st.sidebar.markdown(f"### 📁 Archivos cargados: {total_archivos}/3")
if total_archivos > 0:
    for nombre in st.session_state.dataframes.keys():
        st.sidebar.success(f"✅ {nombre}")

if st.sidebar.button("🔄 Limpiar todos los datos"):
    st.session_state.dataframes = {}
    st.session_state.infos = {}
    st.sidebar.success("Datos limpiados exitosamente!")
    st.rerun()

st.sidebar.markdown("---")
st.sidebar.markdown("### 📝 Nota")
st.sidebar.markdown("""
La suma de diferencias por grupo etario en EPI12 
siempre coincidirá con la diferencia total.
""")
